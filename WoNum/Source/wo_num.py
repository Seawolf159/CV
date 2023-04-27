import os
import pathlib
import shutil
import sys
import time
from configparser import ConfigParser

from docx import Document

import exceptions as exc

if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
    MAIN_DIR = pathlib.Path(sys.executable).parent
    FROZEN_DIR = pathlib.Path(getattr(sys, "_MEIPASS"))
else:
    MAIN_DIR = pathlib.Path(sys.argv[0]).parent
    FROZEN_DIR = pathlib.Path(__file__).parent

FORMS_PATH = MAIN_DIR.joinpath("Forms")
TEMP_PATH = FORMS_PATH.joinpath("ONLY USE ON PRINT FAILURE")


class WoNum:
    def main(self, checked_items, progress, hard_print_amount=1):
        self.remove_temp_directory()
        self.create_temp_directory()

        amount_of_prints = len(checked_items) * hard_print_amount
        sleep_timer = self.read_config()
        sleep_timer = int(sleep_timer["sleep"]["sleep_timer"])
        checked_items.sort(key=lambda x: (x[0], x[1]))
        for selection in checked_items:
            category = selection[0]
            filename = selection[1]

            file_path = FORMS_PATH.joinpath(category.capitalize(), filename)

            for _ in range(hard_print_amount):
                config = self.read_config()
                word_doc = self.open_word_document(file_path)
                next_wo_num = self.get_next_wo_num(config, category)
                text_replacement_mode = self.get_text_replacement_mode(
                    config, category
                )

                if text_replacement_mode == 1:
                    next_page_number = self.get_next_page_number(
                        config, category
                    )
                else:
                    next_page_number = None

                next_wo_num = self.replace_text(
                    word_doc,
                    next_wo_num,
                    text_replacement_mode,
                    filename,
                    next_page_number,
                )

                self.save_word_doc(word_doc, category, filename, next_wo_num)
                self.write_config(
                    config, category, next_wo_num, next_page_number
                )
        self.hard_print(sleep_timer, progress, amount_of_prints)

    @staticmethod
    def create_temp_directory():
        if not TEMP_PATH.exists():
            os.mkdir(TEMP_PATH)

    @staticmethod
    def read_config():
        config = ConfigParser()
        path = MAIN_DIR.joinpath("config.ini")
        config.read(path)
        return config

    @staticmethod
    def open_word_document(file_path):
        word_document = Document(file_path)
        return word_document

    @staticmethod
    def get_next_wo_num(config, category):
        try:
            wo_num = config[category]["num"]
        except KeyError:
            raise exc.ConfigCategoryNotSetUpError(
                f"The '{category}' category hasn't been set-up in the "
                "config.ini file. The config.ini should include this "
                f"section:\n[{category}]\nnum = ....."
            )

        letter = wo_num[0]
        count = int(wo_num[-4:])
        next_num = count + 1
        next_wo_num = f"{letter}{next_num:04}"
        return next_wo_num

    @staticmethod
    def get_text_replacement_mode(config, category):
        try:
            mode = int(config[category]["text_replacement_mode"])
        except KeyError:
            raise exc.ConfigCategoryNotSetUpError(
                f"The '{category}' category hasn't been set-up in the "
                "config.ini file. The config.ini should include this "
                f"section:\n[{category}]\ntext_replacement_mode = ."
            )

        return mode

    @staticmethod
    def get_next_page_number(config, category):
        try:
            page_number = int(config[category]["page_number"])
        except KeyError:
            raise exc.ConfigCategoryNotSetUpError(
                f"The '{category}' category hasn't been set-up in the "
                "config.ini file. The config.ini should include this "
                f"section:\n[{category}]\npage_number = ."
            )
        return page_number + 1

    def replace_text(self, doc, next_wo_num, mode, filename, next_page_number):
        if next_page_number:
            self.replace_footer_text(doc, filename, next_page_number)

        if mode == 0:
            return self.replace_header_text(doc, next_wo_num)
        elif mode == 1:
            return self.replace_table_text(doc, next_wo_num)

    @staticmethod
    def replace_header_text(doc, next_wo_num):
        for paragraph in doc.sections[0].header.paragraphs:
            for run in paragraph.runs:
                if "~" in run.text:
                    run.text = run.text.replace("~", next_wo_num)
        return next_wo_num

    def replace_table_text(self, doc, next_wo_num):
        for row in doc.tables[0].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "~" in run.text:
                            run.text = next_wo_num
                            next_wo_num = self.increment_wo_num(next_wo_num)

        return self.decrement_wo_num(next_wo_num)

    @staticmethod
    def increment_wo_num(wo_num):
        letter = wo_num[0]
        count = int(wo_num[-4:])
        next_num = count + 1
        next_wo_num = f"{letter}{next_num:04}"
        return next_wo_num

    @staticmethod
    def decrement_wo_num(wo_num):
        letter = wo_num[0]
        count = int(wo_num[-4:])
        next_num = count - 1
        next_wo_num = f"{letter}{next_num:04}"
        return next_wo_num

    @staticmethod
    def replace_footer_text(doc, filename, next_page_number):
        for paragraph in doc.sections[0].footer.paragraphs:
            paragraph.text = f"{filename}\t\tPage {next_page_number}"

    @staticmethod
    def save_word_doc(doc, category, filename, next_wo_num):
        part_code = filename.split("-")[0].strip()
        unqiue_directory = TEMP_PATH.joinpath(
            f"{category.capitalize()} {part_code} {next_wo_num}"
        )
        os.mkdir(unqiue_directory)
        doc.save(unqiue_directory.joinpath(filename))

    @staticmethod
    def write_config(config, category, next_wo_num, next_page_number):
        config[category]["num"] = next_wo_num
        if next_page_number:
            config[category]["page_number"] = str(next_page_number)

        path = MAIN_DIR.joinpath("config.ini")
        with open(path, "w") as f:
            config.write(f)

    @staticmethod
    def hard_print(sleep_timer, progress_bar, amount_of_prints):
        progress_increment = 99 / amount_of_prints
        progress_amount = 0
        for doc in TEMP_PATH.glob("*/*.docx"):
            os.startfile(doc, "print")
            time.sleep(sleep_timer)

            progress_amount += progress_increment
            progress_bar.emit(progress_amount)

        time.sleep(sleep_timer + 3)

    @staticmethod
    def remove_temp_directory():
        try:
            if TEMP_PATH.exists():
                shutil.rmtree(TEMP_PATH)
        except PermissionError:
            pass


if __name__ == "__main__":
    pass
